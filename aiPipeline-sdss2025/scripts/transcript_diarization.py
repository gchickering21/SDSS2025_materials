import glob
import os

import nltk
from botocore.exceptions import ClientError
from docx import Document
from docx.shared import RGBColor

import scripts.lib as lib
from genAI_utils.ask_gpt import get_chatgpt_response
from scripts.logging_config import root_logger

nltk.download("punkt")

# load_dotenv()
# # Get the NLTK data path from the .env file
# nltk_data_path = os.getenv('NLTK_DATA')

# if nltk_data_path:
#     nltk.data.path.append(nltk_data_path)
# else:
#     print("Error: NLTK_DATA environment variable not set in the .env file.")


class Diarizer:
    def __init__(self):
        ##TODO: MOVE TO CONFIG
        self.input_transcripts_folder = "word_transcripts"
        self.output_diarized_transcripts_folder = "word_diarized_transcripts"

        self.s3_transcripts_file_input_bucket = "il-empower-transcripts"
        self.s3_transcripts_file_input_folder = "word_transcripts"

    ############################# DOWNLOADING WORD DOCS FROM S3 ###############################

    def download_all_transcripts(self):
        """
        Searches for all files in the S3 input bucket's audio_files folder and adds only valid filenames to the tracker.
        """
        response = self.s3_client.list_objects_v2(
            Bucket=self.s3_transcripts_file_input_bucket,
            Prefix=self.s3_transcripts_file_input_folder + "/",
        )
        if "Contents" in response:
            root_logger.info("Files found in word_transcripts folder:")
            for obj in response["Contents"]:
                try:
                    filename = (
                        obj["Key"].replace(self.input_folder + "/", "").strip()
                    )  # Strip whitespace
                    if filename:  # Check that the filename is not empty
                        s3_key = f"{self.s3_transcripts_file_input_folder}/{filename}"
                        self.s3_client.download_file(
                            self.s3_transcripts_file_input_bucket,
                            s3_key,
                            self.input_transcripts_folder,
                        )
                        root_logger.info(f"Downloaded {filename} from S3")
                except ClientError as e:
                    root_logger.error(f"Failed to download {filename}: {e}")
                    return False

    ############################# UPLOADING DIARIZED WORD DOCS TO S3 ###############################

    def upload_diarized_word_docs_to_s3(self):
        """
        Upload all Word documents in 'word_diarized_transcripts' to the 'transcripts' folder in S3.
        """
        word_files = glob.glob(
            "word_diarized_transcripts/*.docx"
        )  # List all Word files in the new output folder
        for word_file in word_files:
            filename = os.path.basename(word_file)
            s3_key = f"word_diarized_transcripts/{filename}"  # S3 key for the 'transcripts' subfolder
            try:
                # Upload to S3
                self.s3_client.upload_file(word_file, self.s3_transcript_bucket, s3_key)
                root_logger.info(
                    f"Uploaded {filename} to {s3_key} in bucket {self.s3_transcript_bucket}"
                )
            except ClientError as e:
                root_logger.error(f"Failed to upload {filename} to {s3_key}: {e}")
        print("All Word documents have been uploaded to S3.")

    ###################################### PERFORM DIARIZATION #####################################

    def process_and_save_diarized_files(self, file_status=None):
        """
        Loops through all .docx files in the downloaded transcripts folder, applies processing functions,
        and saves the output in the diarized folder with 'diarized_' as a prefix in the filename.
        """
        input_folder = (
            self.input_transcripts_folder
        )  # Folder where files were downloaded
        output_folder = (
            self.output_diarized_transcripts_folder
        )  # Folder to save processed files

        root_logger.info(
            f"starting process for diarization of all files in {input_folder}"
        )
        # Loop through all docx files in the input folder
        try:
            for file_name in os.listdir(input_folder):
                root_logger.info(f"starting process for diarization of {file_name}")
                if file_name.endswith(".docx"):
                    base_filename = os.path.splitext(file_name)[0]
                    # print(base_filename)  # Extract the base filename
                    input_path = os.path.join(input_folder, file_name)
                    # Generate the output filename, replacing "transcript" with "diarized"
                    output_filename = (
                        f"{base_filename.replace('transcribed', 'diarized')}.docx"
                    )
                    output_path = os.path.join(output_folder, output_filename)
                    root_logger.info(f"{output_filename}: about to read in docx file")
                    # Step 1: Extract text from the docx file
                    text = self.read_docx_text(input_path)
                    root_logger.info(f"{output_filename}: about to break into batches")
                    # # Step 2: Break text into batches
                    batches = self.break_into_batches(text)
                    root_logger.info(f"{output_filename} has been broken into batches")
                    # # Step 3: Extract batch dialogue
                    transcribed_name, total_tokens, total_cost = (
                        self.extract_batch_dialogue(batches, output_path)
                    )
                    root_logger.info(
                        f"{output_filename} has been extracted from batch dialogue"
                    )

                    ## STEP 5:
                    lib.clear_folder("diarized_transcribed_text")

                    # Remove the prefix 'transcribed_'
                    without_prefix = file_name.removeprefix("transcribed_")
                    result = without_prefix.split(".")[0]
                    file_status[result]["diarized"] = True
                    # Log success
                    root_logger.info(
                        f"Processed and saved {output_filename} in {output_folder}"
                    )
        except Exception as e:
            root_logger.error(f"An error occurred during diarization process: {e}")

        return file_status

    def read_docx_text(self, docx_path):
        doc = Document(docx_path)
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        return "\n".join(full_text)

    def break_into_batches(self, text_blob, batch_size=100):
        root_logger.info("Running break_into_batches ...")
        sentences = nltk.sent_tokenize(text_blob)
        num_sentences = len(sentences)
        batches = []
        for i in range(0, num_sentences, batch_size):
            batch = sentences[i : i + batch_size]
            batches.append(batch)
        return batches

    def extract_batch_dialogue(self, batches, output_path):
        root_logger.info("Running extract_batch_dialogue ...")
        document = Document()
        counter = 1
        list_of_lists = []
        total_cost = 0
        total_tokens = 0
        for batch in batches:
            # root_logger.info(f'batch number: {batch}')
            TempDocument = Document()
            subset_text = " ".join(batch)
            message_content = []
            # print(f"processing {counter} of {len(batches)}: {len(subset_text.split())} input words")
            prompt = """Review the text between one or more interviewer asking questions to one or more interviewees to learn about their experience.
        Please return the same dialogue, but structured in the following format:
        Interviewer:
        Interviewee:
        ...
        You must include all of the original text."""

            resp, tokens, cost = get_chatgpt_response(
                subset_text,
                sys_context=prompt,
                return_string=False,
                return_token_usage=True,
            )
            total_cost = total_cost + cost
            total_tokens = total_tokens + tokens
            finish_reason = resp.choices[0].finish_reason
            if finish_reason == "stop":
                message_content.append(
                    {
                        "text": resp.choices[0].message.content,
                        "color": RGBColor(0x00, 0x00, 0x00),
                    }
                )
            elif finish_reason == "content_filter":
                message_content.append(
                    {
                        "text": f"""\n--------\nThis section of the transcript was flagged as containing sensitive
                content and was unable to be labeled at this time:\n\n{subset_text}\n--------\n""",
                        "color": RGBColor(0xFF, 0x00, 0x00),
                    }
                )
            elif finish_reason == "length":
                # It could theoretically happen that in the model output, it's adding too much text and causing
                # the token limit to be exceeded (e.g., in cases where it's identifying a lot of switching off between speakers)
                # In such cases, we'll split the batch in half and attempt again.
                # Also if we're hitting this as the finish reason, it means we DIDN'T hit 'content_filter', so no
                # need to have logic for that here.
                split_on = int(len(batch) / 2)
                batch_p1 = " ".join(batch[:split_on])
                batch_p2 = " ".join(batch[split_on:])
                batch_parts = [batch_p1, batch_p2]
                resp_p1, tokens1, cost1 = get_chatgpt_response(
                    batch_p1,
                    sys_context=prompt,
                    return_string=False,
                    return_token_usage=True,
                )
                resp_p2, tokens2, cost2 = get_chatgpt_response(
                    batch_p2,
                    sys_context=prompt,
                    return_string=False,
                    return_token_usage=True,
                )
                cost = cost1 + cost2
                tokens = tokens1 + tokens2
                total_cost = total_cost + cost
                total_tokens = total_tokens + tokens
                responses = [resp_p1, resp_p2]
                for ind, r in enumerate(responses):
                    if r.choices[0].finish_reason == "stop":
                        message_content.append(
                            {
                                "text": r.choices[0].message.content + "\n",
                                "color": RGBColor(0xFF, 0x00, 0x00),
                            }
                        )
                    else:
                        bad_text = batch_parts[ind]
                        error_msg = f"""\n--------\nWe're sorry, OpenAI has encountered an unexpected error and
                        was not able to label this section of the transcript:\n\n{bad_text}\n--------\n"""
                        message_content.append(
                            {"text": error_msg, "color": RGBColor(0xFF, 0x00, 0x00)}
                        )

            elif finish_reason == "null":
                # let's retry
                resp, tokens, cost = get_chatgpt_response(
                    sys_context=prompt, return_string=False, return_token_usage=True
                )
                total_cost = total_cost + cost
                total_tokens = total_tokens + tokens
                finish_reason = resp.choices[0].finish_reason
                if finish_reason != "stop":
                    message_content.append(
                        {
                            "text": f"""\n--------\nWe're sorry, OpenAI has encountered an unexpected error and
                    was not able to label this section of the transcript:\n\n{bad_text}\n--------\n""",
                            "color": RGBColor(0xFF, 0x00, 0x00),
                        }
                    )
                else:
                    message_content.append(
                        {
                            "text": resp.choices[0].message.content,
                            "color": RGBColor(0x00, 0x00, 0x00),
                        }
                    )

            full_content = "\n".join([m["text"] for m in message_content])
            # print(f"\t OUTPUTS: {len(full_content.split())} words; {len(full_content)} characters")
            for content in message_content:
                run = document.add_paragraph().add_run(content["text"])
                font = run.font
                font.color.rgb = content["color"]

                TempDocument.add_paragraph(content["text"])
                TempDocument.save(
                    "diarized_transcribed_text/"
                    + f"test_diarization_{counter}-GPT4-ET-{len(batches)}.docx"
                )
                counter = counter + 1
                list_of_lists.append(
                    [len(subset_text.split()), len(full_content.split())]
                )

        transcribed_name = output_path
        document.save(transcribed_name)
        # df = pd.DataFrame(list_of_lists, columns=["words_sent", "words_recieved"])
        # print(df)
        return transcribed_name, total_tokens, total_cost


###############################################################################################


# def main():
#     diarizer = Diarizer()

#     print("starting second phase")
#     diarizer.process_and_save_diarized_files()


# if __name__ == "__main__":
#     main()
